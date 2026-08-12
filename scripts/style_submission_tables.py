#!/usr/bin/env python3
"""Repair lossy gt DOCX exports for journal submission.

The source DOCX files remain authoritative for all visible strings. This script
changes Word presentation only: page geometry, table widths, rich-text markup,
fills, borders, fonts, alignment, and repeating headers.
"""

from __future__ import annotations

import argparse
import html
import re
from copy import deepcopy
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


FONT_NAME = "Arial"
BLACK = "333333"
WHITE = "FFFFFF"
LIGHT_RULE = "D9D9D9"
MID_RULE = "A6A6A6"
HEADER_GREY = "F3F3F3"
CHEST = "EFC000"
WRIST = "868686"
GLASSES = "0073C2"
CHEST_LIGHT = "FDF9E6"
CHEST_STRONG = "F7DF80"
WRIST_LIGHT = "F3F3F3"
WRIST_STRONG = "C3C3C3"
NEGATIVE_LIGHT = "FFF2F2"
POSITIVE_LIGHT = "E4F5FF"

SITE_COLORS = {
    "Dortmund (DE)": "CC6677",
    "Madrid (ES)": "332288",
    "Izmir (TR)": "88CCEE",
    "Kumasi (GH)": "AA4499",
    "Borås (SE)": "DDCC77",
    "Delft (NL)": "117733",
    "Munich (DE)": "999933",
    "San Pedro, San José (CR)": "44AA99",
}


@dataclass(frozen=True)
class TableJob:
    source: str
    output: str
    kind: str
    title: str


JOBS = (
    TableJob(
        source="Table_metrics_summary.docx",
        output="Table 1.docx",
        kind="summary",
        title="Main Table 1",
    ),
    TableJob(
        source="Table_metrics.docx",
        output="Supplementary Table S5.docx",
        kind="detail",
        title="Supplementary Table S5",
    ),
    TableJob(
        source="Table_metrics_summary_nonwear.docx",
        output="Supplementary Table S6.docx",
        kind="summary",
        title="Supplementary Table S6",
    ),
    TableJob(
        source="Table_metrics_summary_wake_only.docx",
        output="Supplementary Table S7.docx",
        kind="summary",
        title="Supplementary Table S7",
    ),
    TableJob(
        source="Table_crossing.docx",
        output="Supplementary Table S8.docx",
        kind="crossing",
        title="Supplementary Table S8",
    ),
)


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#").upper())


def set_run_font(
    run,
    *,
    size: float,
    color: str = BLACK,
    name: str = FONT_NAME,
    bold: bool | None = None,
    italic: bool | None = None,
    superscript: bool | None = None,
    subscript: bool | None = None,
) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if superscript is not None:
        run.font.superscript = superscript
    if subscript is not None:
        run.font.subscript = subscript


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def format_paragraph(
    paragraph,
    *,
    alignment: WD_ALIGN_PARAGRAPH,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.0,
    keep_with_next: bool = False,
) -> None:
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = keep_with_next


def set_paragraph_border(
    paragraph,
    *,
    side: str,
    color: str,
    size: int,
    space: int = 1,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = p_bdr.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        p_bdr.append(edge)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)


def style_title_paragraph(
    paragraph,
    text: str,
    *,
    size: float,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    after: float = 0,
    top_rule: bool = False,
) -> None:
    clear_paragraph(paragraph)
    paragraph.style = "Normal"
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=False)
    format_paragraph(
        paragraph,
        alignment=alignment,
        after=after,
        line_spacing=1.0,
        keep_with_next=True,
    )
    if top_rule:
        set_paragraph_border(
            paragraph,
            side="top",
            color=MID_RULE,
            size=8,
            space=4,
        )


def replace_svg_with_arrow_tags(markup: str) -> str:
    svg_pattern = re.compile(r"<svg\b.*?</svg>", flags=re.IGNORECASE | re.DOTALL)

    def replace(match: re.Match[str]) -> str:
        svg = match.group(0)
        color_match = re.search(r"fill:(#[0-9A-Fa-f]{6})", svg)
        color = color_match.group(1) if color_match else "#333333"
        if "M214.6 41.4" in svg:
            glyph = "↑"
        elif "M169.4 470.6" in svg:
            glyph = "↓"
        else:
            raise ValueError("Unrecognized Font Awesome arrow path in DOCX export.")
        return f'<arrow color="{color}">{glyph}</arrow>'

    return svg_pattern.sub(replace, markup)


class RichTextWriter(HTMLParser):
    def __init__(self, paragraph, *, size: float, base_color: str, base_bold: bool):
        super().__init__(convert_charrefs=True)
        self.paragraph = paragraph
        self.size = size
        self.styles = [
            {
                "bold": base_bold,
                "italic": False,
                "superscript": False,
                "subscript": False,
                "color": base_color,
            }
        ]

    @property
    def style(self) -> dict[str, object]:
        return self.styles[-1]

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag == "br":
            self.paragraph.add_run().add_break(WD_BREAK.LINE)
            return
        next_style = dict(self.style)
        if tag in {"b", "strong"}:
            next_style["bold"] = True
        elif tag in {"i", "em"}:
            next_style["italic"] = True
        elif tag == "sup":
            next_style["superscript"] = True
            next_style["subscript"] = False
        elif tag == "sub":
            next_style["subscript"] = True
            next_style["superscript"] = False
        elif tag == "arrow":
            next_style["bold"] = True
            next_style["color"] = dict(attrs).get("color", "#333333").lstrip("#")
        elif tag == "color":
            next_style["color"] = dict(attrs).get("value", "#333333").lstrip("#")
        self.styles.append(next_style)

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() != "br" and len(self.styles) > 1:
            self.styles.pop()

    def handle_data(self, data: str) -> None:
        if not data:
            return
        run = self.paragraph.add_run(data)
        set_run_font(
            run,
            size=self.size,
            color=str(self.style["color"]).lstrip("#"),
            bold=bool(self.style["bold"]),
            italic=bool(self.style["italic"]),
            superscript=bool(self.style["superscript"]),
            subscript=bool(self.style["subscript"]),
        )


def write_markup(
    cell,
    markup: str,
    *,
    size: float,
    alignment: WD_ALIGN_PARAGRAPH,
    base_color: str = BLACK,
    base_bold: bool = False,
    line_spacing: float = 1.0,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    format_paragraph(
        paragraph,
        alignment=alignment,
        line_spacing=line_spacing,
    )
    writer = RichTextWriter(
        paragraph,
        size=size,
        base_color=base_color,
        base_bold=base_bold,
    )
    writer.feed(replace_svg_with_arrow_tags(markup))
    writer.close()


def style_existing_cell_text(
    cell,
    *,
    size: float,
    alignment: WD_ALIGN_PARAGRAPH,
    color: str = BLACK,
    bold: bool | None = None,
    line_spacing: float = 1.0,
) -> None:
    for paragraph in cell.paragraphs:
        format_paragraph(
            paragraph,
            alignment=alignment,
            line_spacing=line_spacing,
        )
        for run in paragraph.runs:
            set_run_font(
                run,
                size=size,
                color=color,
                bold=bold if bold is not None else run.bold,
                italic=False,
            )


def unique_cells(row):
    seen: set[int] = set()
    for cell in row.cells:
        identifier = id(cell._tc)
        if identifier not in seen:
            seen.add(identifier)
            yield cell


def set_cell_fill(cell, fill: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)
    if fill is not None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def get_cell_fill(cell) -> str | None:
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return None if shd is None else shd.get(qn("w:fill"))


def set_cell_borders(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_spec in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        if edge_spec is None:
            edge.set(qn("w:val"), "nil")
            for attribute in ("sz", "space", "color"):
                edge.attrib.pop(qn(f"w:{attribute}"), None)
        else:
            color, size = edge_spec
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), str(size))
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), color)


def set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths_inches: list[float]) -> None:
    widths_twips = [round(width * 1440) for width in widths_inches]
    total_twips = sum(widths_twips)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for tag in ("w:tblStyle", "w:tblLook"):
        for node in tbl_pr.findall(qn(tag)):
            tbl_pr.remove(node)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for old_grid in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)

    for row in table.rows:
        grid_index = 0
        for tc in row._tr.tc_lst:
            tc_pr = tc.get_or_add_tcPr()
            grid_span = tc_pr.find(qn("w:gridSpan"))
            span = 1 if grid_span is None else int(grid_span.get(qn("w:val")))
            width = sum(widths_twips[grid_index : grid_index + span])
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            grid_index += span


def set_document_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(9)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_NAME)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_NAME)


def set_page(
    document: Document,
    *,
    landscape: bool,
    margin_inches: float,
    paper_size: str = "A4",
) -> None:
    dimensions = {
        "A4": (210, 297),
        "A3": (297, 420),
    }
    if paper_size not in dimensions:
        raise ValueError(f"Unsupported paper size: {paper_size}.")
    short_edge, long_edge = dimensions[paper_size]
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
        section.page_width = Mm(long_edge if landscape else short_edge)
        section.page_height = Mm(short_edge if landscape else long_edge)
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)
        section.header_distance = Inches(0.1)
        section.footer_distance = Inches(0.1)


def clear_all_cell_presentation(table) -> None:
    for row in table.rows:
        set_row_cant_split(row)
        for cell in unique_cells(row):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_fill(cell, WHITE)
            set_cell_borders(cell, top=None, bottom=None, start=None, end=None)


def source_cell_matrix(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def visible_markup_text(markup: str) -> str:
    markup = replace_svg_with_arrow_tags(markup)
    markup = re.sub(r"<arrow\b[^>]*>", "", markup, flags=re.IGNORECASE)
    markup = re.sub(r"</arrow>", "", markup, flags=re.IGNORECASE)
    markup = re.sub(r"<br\s*/?>", "\n", markup, flags=re.IGNORECASE)
    markup = re.sub(r"<[^>]+>", "", markup)
    return html.unescape(markup)


def normalize_script_digits(text: str) -> str:
    """Normalize glyph-based scripts for comparison with Word run formatting."""
    return text.translate(
        str.maketrans(
            "⁰¹²³⁴⁵⁶⁷⁸⁹₀₁₂₃₄₅₆₇₈₉",
            "01234567890123456789",
        )
    )


def normalize_visible_equivalents(text: str) -> str:
    """Canonicalize visual repairs that replace lossy gt Word fallbacks."""
    text = re.sub(r"\^(-?\d+)", r"\1", text)
    return normalize_script_digits(text.replace("---", "—"))


def assert_table_text_preserved(before: list[list[str]], table) -> None:
    after = source_cell_matrix(table)
    if len(before) != len(after):
        raise AssertionError("Table row count changed during styling.")
    for row_index, (before_row, after_row) in enumerate(zip(before, after, strict=True)):
        if len(before_row) != len(after_row):
            raise AssertionError(f"Table column count changed in row {row_index}.")
        for column_index, (raw, styled) in enumerate(zip(before_row, after_row, strict=True)):
            expected = visible_markup_text(raw)
            if normalize_visible_equivalents(expected) != normalize_visible_equivalents(styled):
                raise AssertionError(
                    "Visible table text changed at "
                    f"row {row_index + 1}, column {column_index + 1}: "
                    f"expected {expected!r}, got {styled!r}."
                )


def footnote_markup(text: str, *, site_colors: bool = False) -> str:
    escaped = html.escape(text)
    marker_match = re.match(r"^([1-9])(.*)$", escaped, flags=re.DOTALL)
    if marker_match:
        escaped = f"<sup>{marker_match.group(1)}</sup>{marker_match.group(2)}"
    escaped = escaped.replace("R²Pos.", "R<sup>2</sup><sub>Pos.</sub>")
    escaped = escaped.replace("R²Id", "R<sup>2</sup><sub>Id</sub>")
    escaped = escaped.replace("R2", "R<sup>2</sup>")
    escaped = escaped.replace("SDId", "SD<sub>Id</sub>")
    if site_colors:
        for label, color in SITE_COLORS.items():
            escaped_label = html.escape(label)
            escaped = escaped.replace(
                escaped_label,
                f'<color value="#{color}"><b>{escaped_label}</b></color>',
            )
    return escaped


def style_summary(document: Document) -> None:
    set_document_defaults(document)
    set_page(document, landscape=False, margin_inches=0.32)
    if len(document.paragraphs) < 5 or len(document.tables) != 1:
        raise ValueError("Unexpected summary-table DOCX structure.")

    title_lines = [
        "Summary overview of outcome-level bias",
        "for body-worn versus eye-level measurement",
    ]
    style_title_paragraph(
        document.paragraphs[0],
        title_lines[0],
        size=13.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        top_rule=True,
    )
    style_title_paragraph(
        document.paragraphs[1],
        title_lines[1],
        size=13.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
    )
    style_title_paragraph(
        document.paragraphs[2],
        document.paragraphs[2].text,
        size=9.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    style_title_paragraph(
        document.paragraphs[3],
        document.paragraphs[3].text,
        size=9.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    style_title_paragraph(
        document.paragraphs[4],
        document.paragraphs[4].text,
        size=9.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=5,
    )
    for paragraph in document.paragraphs[5:]:
        clear_paragraph(paragraph)
        format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    table = document.tables[0]
    before = source_cell_matrix(table)
    if len(table.rows) != 12 or len(table.rows[0].cells) != 4:
        raise ValueError("Unexpected summary-table dimensions.")
    set_table_geometry(table, [1.70, 1.92, 1.92, 2.05])
    clear_all_cell_presentation(table)

    headers = [
        "",
        'Chest<sup>1</sup>',
        'Wrist<sup>1</sup>',
        'General<sup>1,2</sup>',
    ]
    for column, markup in enumerate(headers):
        color = CHEST if column == 1 else WRIST if column == 2 else BLACK
        write_markup(
            table.rows[0].cells[column],
            markup,
            size=9.2,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            base_color=color,
            base_bold=True,
        )
        set_cell_borders(
            table.rows[0].cells[column],
            bottom=(LIGHT_RULE, 8),
            top=None,
            start=None,
            end=None,
        )
        set_cell_margins(
            table.rows[0].cells[column],
            top=40,
            start=55,
            bottom=40,
            end=55,
        )
    set_repeat_table_header(table.rows[0])

    for row_index in range(1, 9):
        row = table.rows[row_index]
        for column, cell in enumerate(row.cells):
            raw = before[row_index][column]
            alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if column in {0, 3}
                else WD_ALIGN_PARAGRAPH.RIGHT
            )
            if column == 0:
                write_markup(
                    cell,
                    html.escape(raw),
                    size=8.2,
                    alignment=alignment,
                    base_bold=True,
                )
            else:
                write_markup(
                    cell,
                    raw,
                    size=8.2,
                    alignment=alignment,
                )
            fill = WHITE
            if column == 1:
                fill = CHEST_STRONG if row_index >= 7 else CHEST_LIGHT
            elif column == 2:
                fill = WRIST_STRONG if row_index >= 7 else WRIST_LIGHT
            set_cell_fill(cell, fill)
            set_cell_margins(cell, top=44, start=55, bottom=44, end=55)
            top_border = ("000000", 14) if row_index == 7 else (LIGHT_RULE, 5)
            set_cell_borders(
                cell,
                top=top_border,
                bottom=None,
                start=(LIGHT_RULE, 5) if column == 1 else None,
                end=None,
            )

    for row_index in range(9, 12):
        cell = next(unique_cells(table.rows[row_index]))
        raw = before[row_index][0]
        markup = (
            f"<b>Note.</b>{html.escape(raw[5:])}"
            if raw.startswith("Note.")
            else footnote_markup(raw)
        )
        write_markup(
            cell,
            markup,
            size=7.0,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=0.95,
        )
        set_cell_fill(cell, WHITE)
        set_cell_margins(cell, top=22, start=45, bottom=20, end=45)
        set_cell_borders(
            cell,
            top=(LIGHT_RULE, 6) if row_index == 9 else None,
            bottom=(MID_RULE, 8) if row_index == 11 else None,
            start=None,
            end=None,
        )

    assert_table_text_preserved(before, table)


def style_crossing(document: Document) -> None:
    set_document_defaults(document)
    set_page(document, landscape=True, margin_inches=0.34)
    if len(document.paragraphs) < 2 or len(document.tables) != 1:
        raise ValueError("Unexpected crossing-table DOCX structure.")

    style_title_paragraph(
        document.paragraphs[0],
        re.sub(r"^Table 1:\s*", "", document.paragraphs[0].text),
        size=14.5,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=3,
        top_rule=True,
    )
    style_title_paragraph(
        document.paragraphs[1],
        document.paragraphs[1].text,
        size=10.0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=5,
    )
    for paragraph in document.paragraphs[2:]:
        clear_paragraph(paragraph)
        format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    table = document.tables[0]
    before = source_cell_matrix(table)
    if len(table.rows) != 10 or len(table.rows[0].cells) != 10:
        raise ValueError("Unexpected crossing-table dimensions.")
    set_table_geometry(table, [2.65, 0.70] + [0.95] * 8)
    clear_all_cell_presentation(table)

    for row_index in range(3):
        for column, cell in enumerate(table.rows[row_index].cells):
            color = CHEST if row_index == 2 and column % 2 == 0 and column >= 2 else BLACK
            if row_index == 2 and column % 2 == 1 and column >= 3:
                color = WRIST
            fill = HEADER_GREY if row_index >= 1 or column >= 2 else WHITE
            set_cell_fill(cell, fill)
            style_existing_cell_text(
                cell,
                size=8.4,
                alignment=(
                    WD_ALIGN_PARAGRAPH.LEFT
                    if row_index == 2 and column == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                ),
                color=color,
                bold=row_index < 2 or column >= 1,
                line_spacing=0.95,
            )
            set_cell_margins(cell, top=35, start=45, bottom=35, end=45)
            set_cell_borders(
                cell,
                top=None,
                bottom=(LIGHT_RULE, 8),
                start=None,
                end=None,
            )
        set_repeat_table_header(table.rows[row_index])

    for row_index in range(3, 9):
        for column, cell in enumerate(table.rows[row_index].cells):
            fill = (
                CHEST_LIGHT
                if column >= 2 and column % 2 == 0
                else WRIST_LIGHT
                if column >= 3 and column % 2 == 1
                else WHITE
            )
            set_cell_fill(cell, fill)
            style_existing_cell_text(
                cell,
                size=9.0,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True if column == 0 else False,
            )
            set_cell_margins(cell, top=48, start=45, bottom=48, end=45)
            set_cell_borders(
                cell,
                top=None,
                bottom=(LIGHT_RULE, 5),
                start=(LIGHT_RULE, 6) if column == 1 else None,
                end=None,
            )

    note_cell = next(unique_cells(table.rows[9]))
    note = before[9][0]
    markup = f"<b>Note.</b>{html.escape(note[5:])}" if note.startswith("Note.") else html.escape(note)
    write_markup(
        note_cell,
        markup,
        size=8.0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=0.95,
    )
    set_cell_fill(note_cell, WHITE)
    set_cell_margins(note_cell, top=35, start=45, bottom=35, end=45)
    set_cell_borders(
        note_cell,
        top=(LIGHT_RULE, 6),
        bottom=(MID_RULE, 8),
        start=None,
        end=None,
    )

    assert_table_text_preserved(before, table)


def build_detailed_header(cell, column: int) -> str:
    labels = {
        0: "",
        1: "",
        2: "Scale<sup>2</sup>",
        3: "Reference<sup>3</sup>",
        4: "",
        5: "Chest<sup>3,4</sup>",
        6: "",
        7: "Wrist<sup>3,4</sup>",
        8: "",
        9: "p<sub>Pos.</sub><sup>5</sup>",
        10: "R<sup>2</sup><sub>Pos.</sub><sup>6</sup>",
        11: "Chest<sup>4</sup>",
        12: "Wrist<sup>4</sup>",
        13: "R<sup>2</sup><sub>Slope</sub><sup>6</sup>",
        14: "Id<sup>4</sup>",
        15: "R<sup>2</sup><sub>Id</sub><sup>6</sup>",
        16: "Date<sup>4</sup>",
        17: "R<sup>2</sup><sub>Date</sub><sup>6</sup>",
        18: "Residual<sup>4</sup>",
        19: "R<sup>2</sup><sub>Res.</sub><sup>6</sup>",
        20: "p<sub>Site</sub><sup>5</sup>",
        21: "p<sub>Inter.</sub><sup>5</sup>",
        22: "R<sup>2</sup><sub>Site</sub><sup>6</sup>",
    }
    return labels[column]


def is_group_row(row) -> bool:
    values = [cell.text for cell in row.cells]
    return bool(values[0]) and len(set(values)) == 1


def style_detail(document: Document) -> None:
    set_document_defaults(document)
    set_page(
        document,
        landscape=True,
        margin_inches=0.20,
        paper_size="A3",
    )
    if len(document.paragraphs) < 3 or len(document.tables) != 1:
        raise ValueError("Unexpected detailed-table DOCX structure.")

    style_title_paragraph(
        document.paragraphs[0],
        re.sub(r"^Table 1:\s*", "", document.paragraphs[0].text),
        size=12.0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        top_rule=True,
    )
    style_title_paragraph(
        document.paragraphs[1],
        document.paragraphs[1].text,
        size=12.0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=2,
    )
    style_title_paragraph(
        document.paragraphs[2],
        document.paragraphs[2].text,
        size=9.0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=3,
    )
    for paragraph in document.paragraphs[3:]:
        clear_paragraph(paragraph)
        format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    table = document.tables[0]
    before = source_cell_matrix(table)
    if len(table.rows) != 68 or len(table.rows[0].cells) != 23:
        raise ValueError("Unexpected detailed-table dimensions.")
    original_fills = [
        [get_cell_fill(cell) for cell in row.cells]
        for row in table.rows
    ]
    set_table_geometry(
        table,
        [
            0.76,
            3.03,
            0.49,
            1.07,
            0.54,
            0.52,
            0.54,
            0.49,
            0.54,
            0.58,
            0.58,
            0.54,
            0.54,
            0.58,
            0.54,
            0.58,
            0.54,
            0.58,
            0.62,
            0.58,
            0.58,
            0.58,
            0.58,
        ],
    )
    clear_all_cell_presentation(table)

    row0_spanners = {
        5: "Chest/wrist versus glasses (estimated bias)",
        11: "Standard deviation (random effect)<sup>1</sup>",
        20: "Site effect",
    }
    for column, cell in enumerate(table.rows[0].cells):
        if column in row0_spanners:
            write_markup(
                cell,
                row0_spanners[column],
                size=9.0,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                base_bold=True,
            )
        elif before[0][column] == "":
            write_markup(
                cell,
                "",
                size=9.0,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
        set_cell_margins(cell, top=22, start=20, bottom=22, end=20)
        set_cell_borders(
            cell,
            top=(MID_RULE, 7),
            bottom=(LIGHT_RULE, 5),
            start=None,
            end=None,
        )
    set_repeat_table_header(table.rows[0])

    for column, cell in enumerate(table.rows[1].cells):
        color = GLASSES if column == 3 else CHEST if column == 5 else WRIST if column == 7 else BLACK
        write_markup(
            cell,
            build_detailed_header(cell, column),
            size=9.0,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            base_color=color,
            base_bold=True,
        )
        set_cell_margins(cell, top=22, start=18, bottom=22, end=18)
        set_cell_borders(
            cell,
            top=None,
            bottom=(LIGHT_RULE, 7),
            start=None,
            end=None,
        )
    set_repeat_table_header(table.rows[1])

    for row_index in range(2, 62):
        row = table.rows[row_index]
        if is_group_row(row):
            cell = next(unique_cells(row))
            style_existing_cell_text(
                cell,
                size=10.0,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
            )
            set_cell_fill(cell, WHITE)
            set_cell_margins(cell, top=26, start=20, bottom=20, end=20)
            set_cell_borders(
                cell,
                top=(LIGHT_RULE, 6),
                bottom=(LIGHT_RULE, 6),
                start=None,
                end=None,
            )
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
                if cell.text == "Timing":
                    paragraph.paragraph_format.page_break_before = True
            continue

        for column, cell in enumerate(row.cells):
            if column in {4, 6, 8}:
                arrow_markup = before[row_index][column].strip()
                write_markup(
                    cell,
                    arrow_markup,
                    size=10.5,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )
            elif before[row_index][column] == "---":
                write_markup(
                    cell,
                    "—",
                    size=10.0,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                )
            elif column == 3 and "^" in before[row_index][column]:
                scientific_markup = re.sub(
                    r"\^(-?\d+)",
                    r"<sup>\1</sup>",
                    html.escape(before[row_index][column]),
                )
                write_markup(
                    cell,
                    scientific_markup,
                    size=10.0,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                )
            else:
                style_existing_cell_text(
                    cell,
                    size=10.0,
                    alignment=(
                        WD_ALIGN_PARAGRAPH.LEFT
                        if column in {0, 1, 2}
                        else WD_ALIGN_PARAGRAPH.RIGHT
                    ),
                    color=BLACK,
                    line_spacing=0.95,
                )
            fill = WHITE
            if original_fills[row_index][column] is not None and column in {5, 7}:
                value = before[row_index][column].strip()
                fill = NEGATIVE_LIGHT if value.startswith(("-", "−")) else POSITIVE_LIGHT
            set_cell_fill(cell, fill)
            set_cell_margins(cell, top=18, start=16, bottom=18, end=16)
            set_cell_borders(
                cell,
                top=None,
                bottom=(LIGHT_RULE, 3),
                start=(LIGHT_RULE, 3) if column in {1, 2} else None,
                end=None,
            )

    for row_index in range(62, 68):
        cell = next(unique_cells(table.rows[row_index]))
        raw = before[row_index][0]
        write_markup(
            cell,
            footnote_markup(raw, site_colors=row_index == 64),
            size=8.5,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=0.95,
        )
        set_cell_fill(cell, WHITE)
        set_cell_margins(cell, top=16, start=20, bottom=14, end=20)
        set_cell_borders(
            cell,
            top=(LIGHT_RULE, 6) if row_index == 62 else None,
            bottom=(MID_RULE, 7) if row_index == 67 else None,
            start=None,
            end=None,
        )

    expected_highlights = sum(
        original_fills[row_index][column] is not None
        for row_index in range(2, 62)
        for column in (5, 7)
    )
    restored_highlights = sum(
        get_cell_fill(table.rows[row_index].cells[column])
        in {NEGATIVE_LIGHT, POSITIVE_LIGHT}
        for row_index in range(2, 62)
        for column in (5, 7)
    )
    if restored_highlights != expected_highlights:
        raise AssertionError(
            "Significance-fill count changed during detailed-table styling: "
            f"expected {expected_highlights}, restored {restored_highlights}."
        )

    assert_table_text_preserved(before, table)


def assert_no_literal_markup(document: Document) -> None:
    forbidden = ("<b>", "<br", "<svg", "</sub>", "</sup>")
    visible_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    ).lower()
    found = [token for token in forbidden if token in visible_text]
    if found:
        raise AssertionError(f"Literal markup remains in styled DOCX: {found}.")


def style_job(source_path: Path, output_path: Path, kind: str) -> None:
    document = Document(source_path)
    if kind == "summary":
        style_summary(document)
    elif kind == "crossing":
        style_crossing(document)
    elif kind == "detail":
        style_detail(document)
    else:
        raise ValueError(f"Unknown table kind: {kind}.")
    assert_no_literal_markup(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    if not output_path.exists() or output_path.stat().st_size <= 0:
        raise IOError(f"Failed to write `{output_path}`.")
    with ZipFile(output_path) as archive:
        corrupt_member = archive.testzip()
    if corrupt_member is not None:
        raise IOError(
            f"Corrupt member `{corrupt_member}` in saved DOCX `{output_path}`."
        )
    reopened = Document(output_path)
    expected_dimensions = {
        "summary": (12, 4),
        "detail": (68, 23),
        "crossing": (10, 10),
    }
    if len(reopened.tables) != 1:
        raise AssertionError(f"Expected one table in `{output_path}`.")
    rows, columns = expected_dimensions[kind]
    saved_table = reopened.tables[0]
    if len(saved_table.rows) != rows or len(saved_table.rows[0].cells) != columns:
        raise AssertionError(
            f"Unexpected saved dimensions in `{output_path}`: "
            f"{len(saved_table.rows)} x {len(saved_table.rows[0].cells)}."
        )
    assert_no_literal_markup(reopened)
    if kind == "detail":
        body_text = "\n".join(
            cell.text
            for row in saved_table.rows[2:62]
            for cell in row.cells
        )
        if "---" in body_text or "^" in body_text:
            raise AssertionError(
                f"Lossy missing-value or exponent fallback remains in `{output_path}`."
            )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source-dir",
        type=Path,
        default=Path("output/tables"),
        help="Directory containing the raw gt DOCX exports.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("manuscript/NatComHealth_R0/Tables"),
        help="Directory for the styled submission-ready DOCX files.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    missing = [job.source for job in JOBS if not (args.source_dir / job.source).exists()]
    if missing:
        raise FileNotFoundError(
            "Missing raw DOCX export(s): " + ", ".join(missing)
        )
    for job in JOBS:
        source = args.source_dir / job.source
        output = args.output_dir / job.output
        style_job(source, output, job.kind)
        print(f"Styled {job.title}: {output}")


if __name__ == "__main__":
    main()
